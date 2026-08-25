"""Deterministic Word export for a structured detailed meeting record."""
from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

BLUE = "174565"
MID_BLUE = "2B6F91"
DARK_BLUE = "234F68"
MUTED = "667784"
LIGHT_BLUE = "EAF3F7"
LIGHT_GRAY = "F2F5F7"
LIGHT_GOLD = "FFF4D6"
LIGHT_RED = "FDEBEC"
WHITE = "FFFFFF"
CONTENT_DXA = 9029
TABLE_INDENT_DXA = 120
# Arial Unicode MS ships with macOS and exposes one Unicode family to Word/WPS
# and LibreOffice, avoiding script-dependent font fallback for Chinese.
EAST_ASIA_FONT = "Arial Unicode MS"
LATIN_FONT = "Arial Unicode MS"


def _shade(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_width(cell: Any, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def _set_table_geometry(table: Any, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            _set_cell_width(cell, widths[min(index, len(widths) - 1)])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = tc_pr.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for edge, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                node = margins.find(qn(f"w:{edge}"))
                if node is None:
                    node = OxmlElement(f"w:{edge}")
                    margins.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")


def _set_repeat_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def _keep_row_together(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _set_run_font(run: Any, *, size: float | None = None, bold: bool | None = None,
                  color: str | None = None, italic: bool | None = None) -> None:
    run.font.name = LATIN_FONT
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, fonts)
    fonts.set(qn("w:ascii"), LATIN_FONT)
    fonts.set(qn("w:hAnsi"), LATIN_FONT)
    fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _format_inline(paragraph: Any, text: str, *, size: float = 10.5,
                   color: str | None = None) -> None:
    parts = re.split(r"(\*\*.*?\*\*|\[S\d{3}\])", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        value = part[2:-2] if bold else part
        ref = bool(re.fullmatch(r"\[S\d{3}\]", value))
        run = paragraph.add_run(value)
        _set_run_font(run, size=size, bold=bold or ref,
                      color=MID_BLUE if ref else color)


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, MID_BLUE, 12, 6),
        ("Heading 3", 11, DARK_BLUE, 8, 4),
    ):
        style = document.styles[name]
        style.font.name = LATIN_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = document.styles[name]
        style.font.name = LATIN_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def _add_field(paragraph: Any, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, text, separate, end])
    _set_run_font(run, size=9, color=MUTED)


def _set_header_footer(document: Document, title: str) -> None:
    section = document.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _format_inline(header_p, f"详细会议记录｜{title}", size=8.5, color=MUTED)
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prefix = footer_p.add_run("第 ")
    _set_run_font(prefix, size=9, color=MUTED)
    _add_field(footer_p, "PAGE")
    middle = footer_p.add_run(" 页 / 共 ")
    _set_run_font(middle, size=9, color=MUTED)
    _add_field(footer_p, "NUMPAGES")
    suffix = footer_p.add_run(" 页")
    _set_run_font(suffix, size=9, color=MUTED)


def _add_callout(document: Document, text: str, fill: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    _set_table_geometry(table, [CONTENT_DXA])
    cell = table.cell(0, 0)
    _shade(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    _format_inline(paragraph, text, size=10.5)


def _parse_markdown_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if cells and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            continue
        rows.append(cells)
    return rows


def _table_widths(rows: list[list[str]]) -> list[int]:
    columns = max(len(row) for row in rows)
    weights = []
    for col in range(columns):
        length = max((len(row[col]) if col < len(row) else 0) for row in rows)
        weights.append(max(6, min(length, 36)))
    total = sum(weights)
    widths = [int(CONTENT_DXA * weight / total) for weight in weights]
    widths[-1] += CONTENT_DXA - sum(widths)
    return widths


def _add_markdown_table(document: Document, lines: list[str]) -> None:
    rows = _parse_markdown_table(lines)
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    _set_table_geometry(table, _table_widths(rows))
    for row_index, values in enumerate(rows):
        for col_index in range(columns):
            cell = table.cell(row_index, col_index)
            value = values[col_index] if col_index < len(values) else ""
            if row_index == 0:
                _shade(cell, BLUE)
            elif row_index % 2 == 0:
                _shade(cell, LIGHT_GRAY)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            _format_inline(paragraph, value, size=9.5,
                           color=WHITE if row_index == 0 else None)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
        if row_index == 0:
            _set_repeat_header(table.rows[0])
        if sum(len(value) for value in values) <= 1000:
            _keep_row_together(table.rows[row_index])


def _add_markdown(document: Document, content: str) -> None:
    lines = content.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if line.strip().startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            _add_markdown_table(document, table_lines)
            continue
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("### "):
            document.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:], level=2)
        elif re.match(r"^[-*]\s+", stripped):
            paragraph = document.add_paragraph(style="List Bullet")
            _format_inline(paragraph, re.sub(r"^[-*]\s+", "", stripped))
        elif re.match(r"^\d+[.、]\s*", stripped):
            paragraph = document.add_paragraph(style="List Number")
            _format_inline(paragraph, re.sub(r"^\d+[.、]\s*", "", stripped))
        else:
            paragraph = document.add_paragraph()
            _format_inline(paragraph, stripped)
        index += 1


def _bookmark(paragraph: Any, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), re.sub(r"[^A-Za-z0-9_]", "_", name)[:38])
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def build_docx(
    record: dict[str, Any],
    output_path: str | Path,
    *,
    source_fragments: list[dict[str, Any]] | None = None,
) -> Path:
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.91)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    _configure_styles(document)
    title = str(record.get("title") or "未命名会议")
    _set_header_footer(document, title)
    document.core_properties.title = title
    document.core_properties.subject = "详细会议记录"
    document.core_properties.author = "会议记录整理系统"

    document.add_paragraph().paragraph_format.space_after = Pt(36)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _format_inline(kicker, "会议记录", size=11, color=MID_BLUE)
    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(18)
    title_p.paragraph_format.space_after = Pt(10)
    title_run = title_p.add_run(title)
    _set_run_font(title_run, size=24, bold=True, color=BLUE)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _format_inline(subtitle, "详细会议记录", size=14, color=DARK_BLUE)

    meta = record.get("meeting_meta") or {}
    meta_rows = [
        ("会议日期", str(meta.get("meeting_date") or "待补充")),
        ("会议背景", str(meta.get("background") or "未填写")),
        ("参会人员", "、".join(
            f"{item.get('name')}（{item.get('role')}）" if item.get("role") else str(item.get("name") or "")
            for item in meta.get("attendees") or []
        ) or "未填写"),
        ("记录状态", f"第 {record.get('revision') or 1} 修订版；根据录音/识别稿整理，待相关人员按需复核"),
    ]
    document.add_paragraph().paragraph_format.space_after = Pt(12)
    table = document.add_table(rows=len(meta_rows), cols=2)
    table.style = "Table Grid"
    _set_table_geometry(table, [1900, CONTENT_DXA - 1900])
    for index, (label, value) in enumerate(meta_rows):
        _shade(table.cell(index, 0), LIGHT_BLUE)
        _format_inline(table.cell(index, 0).paragraphs[0], label, size=9.5, color=BLUE)
        table.cell(index, 0).paragraphs[0].runs[0].bold = True
        _format_inline(table.cell(index, 1).paragraphs[0], value, size=9.5)

    document.add_page_break()
    bookmark_id = 1
    for section_record in record.get("sections") or []:
        heading = document.add_heading(str(section_record.get("title") or "未命名章节"), level=1)
        _bookmark(heading, str(section_record.get("id") or f"section_{bookmark_id}"), bookmark_id)
        bookmark_id += 1
        content = str(section_record.get("content") or "本次会议未形成明确内容")
        section_id = str(section_record.get("id") or "")
        if section_id in {"core-conclusion"}:
            _add_callout(document, content, LIGHT_BLUE)
        elif section_id in {"pending-decisions"}:
            _add_callout(document, content, LIGHT_GOLD)
        elif section_id in {"risks", "recognition-review"} and "|" not in content:
            _add_callout(document, content, LIGHT_RED)
        else:
            _add_markdown(document, content)
        refs = section_record.get("source_refs") or []
        if refs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(8)
            _format_inline(paragraph, f"证据片段：{'、'.join(refs)}", size=8.5, color=MUTED)

    referenced = {
        str(ref)
        for section_record in record.get("sections") or []
        for ref in section_record.get("source_refs") or []
    }
    fragment_rows = [
        fragment for fragment in source_fragments or []
        if str(fragment.get("id") or "") in referenced
    ]
    if fragment_rows:
        heading = document.add_heading("证据片段索引", level=1)
        _bookmark(heading, "source_fragment_index", bookmark_id)
        note = document.add_paragraph()
        _format_inline(
            note,
            "以下摘录用于核对正文中的 S 编号；完整识别稿仍保存在会议应用中，不随 Feishu Archive 同步。",
            size=9,
            color=MUTED,
        )
        table = document.add_table(rows=1 + len(fragment_rows), cols=2)
        table.style = "Table Grid"
        _set_table_geometry(table, [1550, CONTENT_DXA - 1550])
        for index, label in enumerate(("片段", "核对摘录")):
            _shade(table.cell(0, index), BLUE)
            _format_inline(table.cell(0, index).paragraphs[0], label, size=9.5, color=WHITE)
            table.cell(0, index).paragraphs[0].runs[0].bold = True
        _set_repeat_header(table.rows[0])
        _keep_row_together(table.rows[0])
        for row_index, fragment in enumerate(fragment_rows, 1):
            if row_index % 2 == 0:
                _shade(table.cell(row_index, 0), LIGHT_GRAY)
                _shade(table.cell(row_index, 1), LIGHT_GRAY)
            label = f"{fragment.get('id')}\n来源 #{fragment.get('source_id')}"
            excerpt = str(fragment.get("text") or "").strip()
            if len(excerpt) > 800:
                excerpt = excerpt[:800].rstrip() + "……"
            _format_inline(table.cell(row_index, 0).paragraphs[0], label, size=8.5, color=BLUE)
            _format_inline(table.cell(row_index, 1).paragraphs[0], excerpt, size=8.5)
            _keep_row_together(table.rows[row_index])

    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")
    document.save(output)
    return output
