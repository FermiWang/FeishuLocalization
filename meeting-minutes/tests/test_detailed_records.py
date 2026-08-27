import hashlib
import json
import os
import re
import shutil
import subprocess
import tempfile
import threading
import unittest
import zipfile
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from unittest import mock

import httpx
from docx import Document
from fastapi.testclient import TestClient


_DATA_DIR = Path(tempfile.mkdtemp(prefix="meeting-minutes-tests-"))
os.environ["MEETING_MINUTES_DATA_DIR"] = str(_DATA_DIR)

from app import db, llm, processor  # noqa: E402
from app.docx_export import build_docx  # noqa: E402
from app.main import app  # noqa: E402
from app.sources import (  # noqa: E402
    annotate_duplicate_sources,
    effective_sources,
    extract_transcript,
    input_hash,
    stable_fragments,
)


def _structured(title="测试会议"):
    return {
        "title": title,
        "subtitle": "详细会议记录",
        "meeting_meta": {
            "title": title,
            "meeting_date": "2026-08-25",
            "background": "测试",
            "attendees": [{"name": "说话人1", "role": "主持"}],
        },
        "sections": [
            {
                "id": section_id,
                "title": section_title,
                "kind": kind,
                "content": f"{section_title}内容 [S001]",
                "source_refs": ["S001"],
            }
            for section_id, section_title, kind in llm.REQUIRED_SECTIONS
        ],
        "provenance": {
            "model_id": llm.EXACT_MODEL_ID,
            "prompt_version": llm.PROMPT_VERSION,
        },
    }


class DetailedRecordDatabaseTests(unittest.TestCase):
    def setUp(self):
        db.init_db()
        con = db.get_conn()
        for table in (
            "sync_events", "record_revisions", "processing_jobs", "meeting_sources",
            "minutes", "attendees", "meetings",
        ):
            con.execute(f"DELETE FROM {table}")
        con.commit()

    def test_multi_source_order_pair_revision_conflict_and_append_only_events(self):
        meeting_id = db.create_meeting("测试会议", "2026-08-25", "背景", [])
        first_path = db.UPLOAD_DIR / "first.txt"
        second_path = db.UPLOAD_DIR / "second.txt"
        first_path.write_text("第一段", encoding="utf-8")
        second_path.write_text("第二段", encoding="utf-8")
        digest = hashlib.sha256("同一哈希".encode()).hexdigest()
        first = db.add_source(
            meeting_id, source_type="transcript", original_name="first.txt",
            stored_path=str(first_path), sha256=digest, text_content="第一段",
        )
        second = db.add_source(
            meeting_id, source_type="transcript", original_name="second.txt",
            stored_path=str(second_path), sha256=digest, text_content="第二段",
        )
        db.reorder_sources(meeting_id, [second["id"], first["id"]])
        db.set_source_pair(meeting_id, second["id"], "pair-A")
        ordered = db.list_sources(meeting_id)
        self.assertEqual([item["id"] for item in ordered], [second["id"], first["id"]])
        self.assertEqual(ordered[0]["pair_key"], "pair-A")

        meeting = db.get_meeting(meeting_id)
        saved = db.save_revision(
            meeting_id, _structured(), "# 测试会议\n", input_hash=input_hash(meeting),
            model_id=llm.EXACT_MODEL_ID, prompt_version=llm.PROMPT_VERSION,
            editor_kind="model", base_revision=0,
            source_fragments=[{
                "id": "S001", "source_id": first["id"],
                "text": "第一段", "speaker_labeled": False,
            }],
        )
        self.assertEqual(saved["revision"], 1)
        with self.assertRaisesRegex(ValueError, "revision conflict"):
            db.save_revision(
                meeting_id, _structured(), "# 冲突\n", input_hash=input_hash(meeting),
                model_id=llm.EXACT_MODEL_ID, prompt_version=llm.PROMPT_VERSION,
                editor_kind="user", base_revision=0,
            )
        second_revision = db.save_revision(
            meeting_id, _structured(), "# 修订\n", input_hash=input_hash(meeting),
            model_id=llm.EXACT_MODEL_ID, prompt_version=llm.PROMPT_VERSION,
            editor_kind="user", base_revision=1,
        )
        self.assertEqual(second_revision["revision"], 2)
        self.assertEqual(second_revision["source_fragments"][0]["id"], "S001")
        events = db.export_events(0)["events"]
        self.assertEqual([item["revision"] for item in events], [1, 2])
        self.assertNotIn("stored_path", json.dumps(events, ensure_ascii=False))
        self.assertNotIn("第一段", json.dumps(events, ensure_ascii=False))

        files = db.delete_meeting(meeting_id)
        self.assertIn(str(first_path), files)
        tombstone = db.export_events(events[-1]["seq"])["events"]
        self.assertEqual(tombstone[0]["event_type"], "delete")

    def test_job_recovery_and_input_hash_reset(self):
        meeting_id = db.create_meeting("测试会议", "2026-08-25", "", [])
        path = db.UPLOAD_DIR / "input.txt"
        path.write_text("内容", encoding="utf-8")
        db.add_source(
            meeting_id, source_type="transcript", original_name="input.txt",
            stored_path=str(path), sha256=db.file_sha256(path), text_content="内容",
        )
        meeting = db.get_meeting(meeting_id)
        job = db.enqueue_job(meeting_id, input_hash(meeting))
        claimed = db.claim_next_job()
        self.assertEqual(claimed["id"], job["id"])
        db.update_job(job["id"], checkpoint={"extracted": [{"x": 1}]})
        self.assertEqual(db.recover_jobs(), 1)
        db.reset_job_input(job["id"], "new-hash", status="queued")
        reset = db.get_job(job["id"])
        self.assertEqual(reset["input_hash"], "new-hash")
        self.assertEqual(json.loads(reset["checkpoint_json"]), {})

    def test_failed_job_with_same_input_reuses_saved_checkpoint(self):
        meeting_id = db.create_meeting("断点复用", "2026-08-25", "", [])
        job = db.enqueue_job(meeting_id, "same-input")
        claimed = db.claim_next_job()
        self.assertEqual(claimed["id"], job["id"])
        db.update_job(
            job["id"],
            checkpoint={"fragment_ids": ["S001"], "extracted": [{"items": []}]},
        )
        db.finish_job(job["id"], error="Expecting ',' delimiter")
        resumed = db.enqueue_job(meeting_id, "same-input")
        self.assertEqual(resumed["id"], job["id"])
        self.assertEqual(resumed["status"], "queued")
        self.assertEqual(resumed["error"], "")
        self.assertEqual(len(json.loads(resumed["checkpoint_json"])["extracted"]), 1)

    def test_parallel_claims_are_atomic_and_unique(self):
        for index in range(6):
            meeting_id = db.create_meeting(f"并发会议{index}", "2026-08-25", "", [])
            db.enqueue_job(meeting_id, f"hash-{index}")
        with ThreadPoolExecutor(max_workers=6) as pool:
            claimed = list(pool.map(lambda _: db.claim_next_job(), range(6)))
        ids = [job["id"] for job in claimed if job]
        self.assertEqual(len(ids), 6)
        self.assertEqual(len(set(ids)), 6)

    def test_completed_legacy_six_section_result_migrates_without_rewriting_original(self):
        meeting_id = db.create_meeting("历史会议", "2026-08-20", "旧数据", [])
        original = (
            "## 会议纪要\n讨论消防水压。\n\n"
            "## 会议达成一致\n先开展试点。\n\n"
            "## 会议未达成一致\n传感器型号待定。\n\n"
            "## 会议待办\n- [ ] 完成 POC"
        )
        db.set_minutes_status(meeting_id, "done", content=original)
        con = db.get_conn()
        con.execute("DELETE FROM schema_meta WHERE key='legacy_minutes_v1'")
        con.commit()
        db._migrate_legacy_minutes()
        con.commit()
        record = db.get_record(meeting_id)
        self.assertEqual(record["revision"], 1)
        self.assertEqual(record["markdown"], original)
        self.assertEqual(record["editor_kind"], "legacy-migration")
        self.assertIn(
            "不补造证据引用",
            next(
                section["content"] for section in record["structured"]["sections"]
                if section["id"] == "compilation-notes"
            ),
        )
        self.assertEqual(db.export_events(0)["events"][0]["event_type"], "upsert")

        interrupted_id = db.create_meeting("中断会议", "2026-08-20", "", [])
        db.set_minutes_status(interrupted_id, "processing")
        db._recover_legacy_processing_state()
        con.commit()
        interrupted = db.get_minutes(interrupted_id)
        self.assertEqual(interrupted["status"], "pending")
        self.assertIn("重新开始整理", interrupted["error"])


class SourceAndDocumentTests(unittest.TestCase):
    def test_docx_and_txt_are_extracted_and_fragments_have_stable_ids(self):
        source = _DATA_DIR / "识别稿.docx"
        document = Document()
        document.add_paragraph("第一段会议内容。")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "行动"
        table.cell(0, 1).text = "开展测试"
        document.save(source)
        text = extract_transcript(source)
        self.assertIn("第一段会议内容", text)
        self.assertIn("行动 | 开展测试", text)
        fragments = stable_fragments([
            {"source_id": 5, "position": 0, "text": text, "speaker_labeled": False},
        ], max_chars=12)
        self.assertEqual(fragments[0]["id"], "S001")
        self.assertEqual(
            [item["id"] for item in fragments],
            [f"S{index:03d}" for index in range(1, len(fragments) + 1)],
        )

    def test_exact_duplicate_sources_are_annotated_and_excluded_from_effective_input(self):
        sources = [
            {
                "id": 10, "source_type": "transcript", "position": 0,
                "pair_key": "", "sha256": "same", "generated": 0,
            },
            {
                "id": 11, "source_type": "transcript", "position": 1,
                "pair_key": "", "sha256": "same", "generated": 0,
            },
            {
                "id": 12, "source_type": "transcript", "position": 2,
                "pair_key": "other-audio", "sha256": "same", "generated": 0,
            },
        ]
        annotated = annotate_duplicate_sources(sources)
        self.assertIsNone(annotated[0]["duplicate_of_source_id"])
        self.assertEqual(annotated[1]["duplicate_of_source_id"], 10)
        self.assertIsNone(annotated[2]["duplicate_of_source_id"])
        self.assertEqual([item["id"] for item in effective_sources(sources)], [10, 12])

    def test_legacy_duplicate_checkpoint_reuses_only_the_proven_prefix(self):
        fragments = [
            {"id": "S001", "text": "第一段"},
            {"id": "S002", "text": "第二段"},
        ]
        checkpoint = {
            "fragment_ids": ["S001", "S002", "S003", "S004"],
            "extracted": [{"items": [1]}, {"items": [2]}, {"items": [3]}, {"items": [4]}],
        }
        reconciled = llm.reconcile_checkpoint(checkpoint, fragments)
        self.assertEqual(len(reconciled["extracted"]), 2)
        self.assertEqual(reconciled["fragment_ids"], ["S001", "S002"])
        self.assertIn("复用", reconciled["checkpoint_note"])

    def test_word_export_contains_all_sections_tables_and_page_fields(self):
        output = _DATA_DIR / "record.docx"
        record = _structured()
        record["sections"][2]["content"] = "|议题|结论|\n|---|---|\n|消防|开展 POC [S001]|"
        build_docx(
            record,
            output,
            source_fragments=[{
                "id": "S001", "source_id": 9,
                "text": "会议原始片段用于核对。", "speaker_labeled": False,
            }],
        )
        self.assertTrue(output.is_file())
        with zipfile.ZipFile(output) as package:
            document_xml = package.read("word/document.xml").decode("utf-8")
            footer_xml = package.read("word/footer1.xml").decode("utf-8")
        self.assertIn("核心结论", document_xml)
        self.assertIn("转写辨识与复核清单", document_xml)
        self.assertIn("证据片段索引", document_xml)
        self.assertIn("tblHeader", document_xml)
        self.assertIn("PAGE", footer_xml)
        self.assertIn("NUMPAGES", footer_xml)


class ProcessorInputMatrixTests(unittest.TestCase):
    def setUp(self):
        db.init_db()
        con = db.get_conn()
        for table in (
            "sync_events", "record_revisions", "processing_jobs", "meeting_sources",
            "minutes", "attendees", "meetings",
        ):
            con.execute(f"DELETE FROM {table}")
        con.commit()

    def _source(self, meeting_id, name, source_type, text="", pair_key=""):
        path = db.UPLOAD_DIR / name
        path.write_bytes(text.encode("utf-8") if source_type == "transcript" else b"RIFFaudio")
        return db.add_source(
            meeting_id,
            source_type=source_type,
            original_name=name,
            stored_path=str(path),
            sha256=db.file_sha256(path),
            text_content=text,
            pair_key=pair_key,
        )

    def test_authoritative_transcript_is_aligned_while_unpaired_audio_is_transcribed(self):
        meeting_id = db.create_meeting("混合来源", "2026-08-25", "", [])
        transcript = self._source(
            meeting_id, "authoritative.txt", "transcript", "权威识别稿原文。", "pair-A"
        )
        paired_audio = self._source(meeting_id, "paired.wav", "audio", pair_key="pair-A")
        unpaired_audio = self._source(meeting_id, "unpaired.wav", "audio")
        meeting = db.get_meeting(meeting_id)
        job = db.enqueue_job(meeting_id, input_hash(meeting))
        authoritative_inputs = []

        def analyze(_meeting_id, source_id, _path, authoritative_text, _upload_dir):
            authoritative_inputs.append((source_id, authoritative_text))
            if authoritative_text:
                return {
                    "text": authoritative_text,
                    "labeled": "发言人2：权威识别稿原文。",
                    "segments": [{"speaker": 1, "start_ms": 0, "end_ms": 1000}],
                }
            return {
                "text": "自动转写内容。",
                "labeled": "发言人1：自动转写内容。",
                "segments": [{"speaker": 0, "start_ms": 0, "end_ms": 800}],
            }

        with mock.patch("app.processor.spk.available", return_value=True), mock.patch(
            "app.processor.spk.analyze", side_effect=analyze
        ):
            fragments, policy = processor.prepare_fragments(job, meeting)
        self.assertEqual(authoritative_inputs[0], (paired_audio["id"], "权威识别稿原文。"))
        self.assertEqual(authoritative_inputs[1], (unpaired_audio["id"], None))
        self.assertTrue(policy["transcript_authoritative"])
        self.assertTrue(policy["audio_alignment_used"])
        self.assertTrue(policy["audio_asr_used"])
        combined = "\n".join(item["text"] for item in fragments)
        self.assertIn("发言人2：权威识别稿原文", combined)
        self.assertIn("发言人1：自动转写内容", combined)
        generated = [item for item in db.list_sources(meeting_id) if item["generated"]]
        self.assertEqual(len(generated), 1)
        self.assertEqual(generated[0]["pair_key"], f"audio:{unpaired_audio['id']}")
        self.assertEqual(transcript["text_content"], "权威识别稿原文。")

    def test_audio_only_fails_visibly_when_funasr_is_unavailable(self):
        meeting_id = db.create_meeting("仅录音", "2026-08-25", "", [])
        self._source(meeting_id, "only.wav", "audio")
        meeting = db.get_meeting(meeting_id)
        job = db.enqueue_job(meeting_id, input_hash(meeting))
        with mock.patch("app.processor.spk.available", return_value=False):
            with self.assertRaisesRegex(RuntimeError, "FunASR 环境不可用"):
                processor.prepare_fragments(job, meeting)

    def test_paired_audio_does_not_block_authoritative_transcript_without_funasr(self):
        meeting_id = db.create_meeting("识别稿优先", "2026-08-25", "", [])
        self._source(
            meeting_id, "authority.txt", "transcript", "权威文本保持不变。", "pair-B"
        )
        self._source(meeting_id, "auxiliary.wav", "audio", pair_key="pair-B")
        meeting = db.get_meeting(meeting_id)
        job = db.enqueue_job(meeting_id, input_hash(meeting))
        with mock.patch("app.processor.spk.available", return_value=False):
            fragments, policy = processor.prepare_fragments(job, meeting)
        self.assertEqual(len(fragments), 1)
        self.assertIn("权威文本保持不变", fragments[0]["text"])
        self.assertTrue(policy["transcript_authoritative"])
        self.assertFalse(policy["audio_asr_used"])

    def test_duplicate_unpaired_transcript_is_skipped_without_deleting_raw_source(self):
        meeting_id = db.create_meeting("重复识别稿", "2026-08-25", "", [])
        first = self._source(meeting_id, "same-1.txt", "transcript", "相同内容。")
        second = self._source(meeting_id, "same-2.txt", "transcript", "相同内容。")
        meeting = db.get_meeting(meeting_id)
        job = db.enqueue_job(meeting_id, input_hash(meeting))
        fragments, policy = processor.prepare_fragments(job, meeting)
        self.assertEqual(len(fragments), 1)
        self.assertEqual(fragments[0]["source_id"], first["id"])
        self.assertEqual(policy["duplicate_sources_skipped"], [second["id"]])
        self.assertEqual(len(db.list_sources(meeting_id)), 2)


class ExactModelTests(unittest.TestCase):
    @staticmethod
    def _response(value):
        return httpx.Response(200, json=value, request=httpx.Request("GET", "http://local"))

    @staticmethod
    def _text_response(value):
        return httpx.Response(
            200, text=value, request=httpx.Request("GET", "http://local")
        )

    def test_preflight_requires_exact_model_and_honors_vllm_capacity(self):
        health = self._text_response("")
        models = self._response({
            "data": [{"id": llm.EXACT_MODEL_ID}, {"id": "Qwen3.6-27B-FP8"}]
        })
        idle_metrics = self._text_response(
            '# TYPE vllm:num_requests_running gauge\n'
            f'vllm:num_requests_running{{engine="0",model_name="{llm.EXACT_MODEL_ID}"}} 0.0\n'
            '# TYPE vllm:num_requests_waiting gauge\n'
            f'vllm:num_requests_waiting{{engine="0",model_name="{llm.EXACT_MODEL_ID}"}} 0.0\n'
        )
        with mock.patch(
            "app.llm.httpx.get", side_effect=[health, models, idle_metrics]
        ):
            self.assertEqual(llm.model_preflight()["model_id"], llm.EXACT_MODEL_ID)
        available_metrics = self._text_response(
            f'vllm:num_requests_running{{engine="0",model_name="{llm.EXACT_MODEL_ID}"}} 1.0\n'
            f'vllm:num_requests_waiting{{engine="0",model_name="{llm.EXACT_MODEL_ID}"}} 0.0\n'
        )
        with mock.patch(
            "app.llm.httpx.get", side_effect=[health, models, available_metrics]
        ):
            result = llm.model_preflight()
            self.assertEqual(result["scheduler"]["num_running"], 1)
        full_but_prioritizable_metrics = self._text_response(
            f'vllm:num_requests_running{{engine="0",model_name="{llm.EXACT_MODEL_ID}"}} 7.0\n'
            f'vllm:num_requests_waiting{{engine="0",model_name="{llm.EXACT_MODEL_ID}"}} 1.0\n'
        )
        with mock.patch(
            "app.llm.httpx.get",
            side_effect=[health, models, full_but_prioritizable_metrics],
        ):
            result = llm.model_preflight()
            self.assertEqual(result["request_priority"], -(2**63))
        congested_metrics = self._text_response(
            f'vllm:num_requests_running{{engine="0",model_name="{llm.EXACT_MODEL_ID}"}} 8.0\n'
            f'vllm:num_requests_waiting{{engine="0",model_name="{llm.EXACT_MODEL_ID}"}} '
            f'{float(llm.MODEL_MAX_WAITING_REQUESTS)}\n'
        )
        with mock.patch(
            "app.llm.httpx.get", side_effect=[health, models, congested_metrics]
        ):
            with self.assertRaises(llm.ModelBusyError):
                llm.model_preflight()
        aliases_only = self._response({"data": [{"id": "Qwen3.6-27B-FP8"}]})
        with mock.patch("app.llm.httpx.get", side_effect=[health, aliases_only]):
            with self.assertRaises(llm.ModelIdentityError):
                llm.model_preflight()
        with mock.patch.object(llm, "CONFIGURED_MODEL", "alias-model"):
            with self.assertRaises(llm.ModelIdentityError):
                llm.model_preflight()

    def test_invalid_inline_reference_is_rejected(self):
        value = _structured()
        value["sections"][0]["content"] = "虚构内容 [S999]"
        with self.assertRaisesRegex(ValueError, "虚构片段引用"):
            llm._normalize_final(value, {"title": "测试"}, {"S001"})

    def test_extract_json_repairs_bounded_missing_comma(self):
        repaired = llm._extract_json('{"first": 1\n"second": [2]\n"third": {"x": 3}}')
        self.assertEqual(repaired, {"first": 1, "second": [2], "third": {"x": 3}})

    def test_repeated_invalid_json_is_recoverable_and_repair_sees_full_output(self):
        long_invalid = "x" * 20_000
        with mock.patch(
            "app.llm._chat",
            side_effect=[
                llm.ChatResult(long_invalid, "stop", {}, "first"),
                llm.ChatResult("still invalid", "stop", {}, "repair"),
            ],
        ) as chat:
            with self.assertRaises(llm.ModelDeterministicError):
                llm._chat_json([{"role": "user", "content": "test"}], 100)
        repair_messages = chat.call_args_list[1].args[0]
        self.assertEqual(len(repair_messages[-2]["content"]), len(long_invalid))

    def test_preflight_network_timeout_is_recoverable(self):
        with mock.patch("app.llm.httpx.get", side_effect=httpx.ConnectTimeout("timeout")):
            with self.assertRaises(llm.ModelTemporaryError):
                llm.model_preflight()

    def test_chat_streams_and_sends_only_exact_model(self):
        response = httpx.Response(
            200,
            text=(
                'data: {"choices":[{"delta":{"content":"{"}}]}\n\n'
                'data: {"choices":[{"delta":{"content":"}"},"finish_reason":"stop"}]}\n\n'
                'data: [DONE]\n\n'
            ),
            request=httpx.Request("POST", "http://local"),
        )
        context = mock.MagicMock()
        context.__enter__.return_value = response
        with mock.patch("app.llm.model_preflight"), mock.patch(
            "app.llm.httpx.stream", return_value=context,
        ) as stream:
            result = llm._chat([{"role": "user", "content": "x"}], 50)
            self.assertEqual(result.text, "{}")
            self.assertEqual(result.finish_reason, "stop")
        self.assertEqual(stream.call_count, 1)
        payload = stream.call_args.kwargs["json"]
        self.assertEqual(payload["model"], llm.EXACT_MODEL_ID)
        self.assertEqual(payload["max_completion_tokens"], 50)
        self.assertEqual(payload["reasoning_effort"], "none")
        self.assertEqual(payload["thinking_token_budget"], 0)
        self.assertIs(payload["include_reasoning"], False)
        self.assertEqual(payload["chat_template_kwargs"], {"enable_thinking": False})
        self.assertEqual(payload["response_format"], {"type": "json_object"})
        self.assertEqual(payload["priority"], -(2**63))
        self.assertIs(payload["stream"], True)
        self.assertEqual(
            stream.call_args.kwargs["headers"]["X-Vllm-Priority"], str(-(2**63))
        )

    def test_chat_timeout_is_recoverable_and_not_immediately_duplicated(self):
        response = mock.MagicMock()
        response.raise_for_status.return_value = None
        response.iter_lines.side_effect = httpx.ReadTimeout("timeout")
        context = mock.MagicMock()
        context.__enter__.return_value = response
        with mock.patch("app.llm.model_preflight"), mock.patch(
            "app.llm.httpx.stream", return_value=context,
        ) as stream:
            with self.assertRaises(llm.ModelTemporaryError):
                llm._chat([{"role": "user", "content": "x"}], 50)
        self.assertEqual(stream.call_count, 1)

    def test_chat_length_finish_is_detected_before_json_repair(self):
        response = httpx.Response(
            200,
            text=(
                'data: {"choices":[{"delta":{"content":"{"}}]}\n\n'
                'data: {"choices":[{"delta":{},"finish_reason":"length"}]}\n\n'
                'data: [DONE]\n\n'
            ),
            request=httpx.Request("POST", "http://local"),
        )
        context = mock.MagicMock()
        context.__enter__.return_value = response
        with mock.patch("app.llm.model_preflight"), mock.patch(
            "app.llm.httpx.stream", return_value=context,
        ):
            with self.assertRaises(llm.ModelOutputTruncatedError):
                llm._chat_json(
                    [{"role": "user", "content": "x"}], 1,
                    response_schema=llm.EXTRACTION_SCHEMA,
                    schema_name="length-test",
                )

    def test_aggregation_length_failure_splits_and_caches_each_half(self):
        first = {"topics": [], "items": [{
            "kind": "meeting_fact", "text": "一", "speaker": "", "owner": "",
            "deadline": "", "source_refs": ["S001"],
        }]}
        second = {"topics": [], "items": [{
            "kind": "meeting_fact", "text": "二", "speaker": "", "owner": "",
            "deadline": "", "source_refs": ["S002"],
        }]}
        truncated = llm.ModelOutputTruncatedError(
            "too long", code="output_truncated", fingerprint="root",
        )
        cache = {}
        with mock.patch(
            "app.llm._chat_json", side_effect=[truncated, first, second],
        ) as chat:
            result = llm._aggregate_adaptive(
                [first, second], {"S001", "S002"}, cache,
                label="议题聚合 1/1", on_progress=None, progress=70,
            )
        self.assertEqual(result, [first, second])
        self.assertEqual(chat.call_count, 3)
        self.assertEqual(len(cache), 3)

        with mock.patch("app.llm._chat_json") as resumed_chat:
            resumed = llm._aggregate_adaptive(
                [first, second], {"S001", "S002"}, cache,
                label="议题聚合 1/1", on_progress=None, progress=70,
            )
        self.assertEqual(resumed, [first, second])
        resumed_chat.assert_not_called()

    def test_final_batch_length_failure_split_is_reused_after_resume(self):
        specifications = llm.REQUIRED_SECTIONS[:2]

        def response_for(specification):
            section_id, title, kind = specification
            return {
                "title": "测试会议",
                "sections": [{
                    "id": section_id, "title": title, "kind": kind,
                    "content": "有证据的内容 [S001]", "source_refs": ["S001"],
                }],
                "recognition_notes": [],
            }

        truncated = llm.ModelOutputTruncatedError(
            "too long", code="output_truncated", fingerprint="final-root",
        )
        cache = {}
        with mock.patch(
            "app.llm._chat_json",
            side_effect=[
                truncated,
                response_for(specifications[0]),
                response_for(specifications[1]),
            ],
        ) as chat:
            sections, notes = llm._generate_sections_adaptive(
                specifications,
                {"evidence": "[]", "valid_source_refs": ["S001"]},
                {"S001"}, cache, on_progress=None, completed_before=0,
            )
        self.assertEqual([section["id"] for section in sections], [item[0] for item in specifications])
        self.assertEqual(notes, [])
        self.assertEqual(chat.call_count, 3)
        self.assertEqual(len(cache), 3)

        with mock.patch("app.llm._chat_json") as resumed_chat:
            resumed, _ = llm._generate_sections_adaptive(
                specifications,
                {"evidence": "[]", "valid_source_refs": ["S001"]},
                {"S001"}, cache, on_progress=None, completed_before=0,
            )
        self.assertEqual([section["id"] for section in resumed], [item[0] for item in specifications])
        resumed_chat.assert_not_called()


class ParallelProcessorTests(unittest.TestCase):
    def setUp(self):
        db.init_db()
        con = db.get_conn()
        for table in (
            "sync_events", "record_revisions", "processing_jobs", "meeting_sources",
            "minutes", "attendees", "meetings",
        ):
            con.execute(f"DELETE FROM {table}")
        con.commit()

    def test_worker_pool_starts_bounded_parallel_workers_only_once(self):
        processor.stop_worker()
        ready = threading.Barrier(4)
        release = threading.Event()

        def held_loop():
            ready.wait(timeout=2)
            release.wait(timeout=2)

        try:
            with mock.patch.object(processor, "MAX_PARALLEL_JOBS", 3), mock.patch(
                "app.processor.db.recover_jobs"
            ) as recover, mock.patch("app.processor._loop", side_effect=held_loop):
                processor.start_worker()
                ready.wait(timeout=2)
                self.assertEqual(processor.worker_status()["active_workers"], 3)
                processor.start_worker()
                self.assertEqual(recover.call_count, 1)
        finally:
            release.set()
            processor.stop_worker()

    def test_temporary_model_failure_requeues_without_losing_checkpoint(self):
        meeting_id = db.create_meeting("临时错误恢复", "2026-08-25", "", [])
        source_path = db.UPLOAD_DIR / "temporary-retry.txt"
        source_path.write_text("测试内容", encoding="utf-8")
        db.add_source(
            meeting_id, source_type="transcript", original_name="temporary-retry.txt",
            stored_path=str(source_path), sha256="hash", text_content="测试内容",
        )
        meeting = db.get_meeting(meeting_id)
        job = db.enqueue_job(meeting_id, input_hash(meeting))
        db.update_job(
            job["id"],
            checkpoint={"fragment_ids": ["S001"], "extracted": [{"x": 1}]},
        )
        claimed = db.claim_next_job()
        with mock.patch(
            "app.processor.prepare_fragments",
            return_value=([{"id": "S001", "text": "测试内容"}], {}),
        ), mock.patch(
            "app.processor.llm.organize",
            side_effect=llm.ModelTemporaryError(
                "模型流式连接超时或中断，已保留断点等待重试"
            ),
        ), mock.patch.object(processor._stop, "wait", return_value=False):
            processor.process_job(claimed)
        recovered = db.get_job(job["id"])
        self.assertEqual(recovered["status"], "queued")
        self.assertIn("extracted", recovered["checkpoint_json"])
        self.assertEqual(db.get_minutes(meeting_id)["status"], "processing")

    def test_same_deterministic_failure_stops_after_three_attempts(self):
        meeting_id = db.create_meeting("确定性失败限次", "2026-08-25", "", [])
        source_path = db.UPLOAD_DIR / "deterministic.txt"
        source_path.write_text("测试内容", encoding="utf-8")
        db.add_source(
            meeting_id, source_type="transcript", original_name="deterministic.txt",
            stored_path=str(source_path), sha256=db.file_sha256(source_path),
            text_content="测试内容",
        )
        meeting = db.get_meeting(meeting_id)
        db.enqueue_job(meeting_id, input_hash(meeting))
        failure = llm.ModelDeterministicError(
            "相同输出失败", code="invalid_json", fingerprint="same-request",
        )
        with mock.patch(
            "app.processor.prepare_fragments",
            return_value=([{"id": "S001", "text": "测试内容"}], {}),
        ), mock.patch(
            "app.processor.llm.organize", side_effect=failure,
        ), mock.patch.object(processor._stop, "wait", return_value=False):
            for _ in range(3):
                claimed = db.claim_next_job()
                self.assertIsNotNone(claimed)
                processor.process_job(claimed)
        finished = db.get_latest_job(meeting_id)
        self.assertEqual(finished["status"], "failed")
        self.assertEqual(finished["same_failure_count"], 3)
        self.assertEqual(finished["last_error_code"], "invalid_json")
        self.assertIn("停止自动重试", finished["error"])


class MeetingApiTests(unittest.TestCase):
    def test_upload_response_does_not_echo_transcript_body(self):
        headers = {"X-Meeting-Minutes-Action": "confirm"}
        with TestClient(app) as client:
            created = client.post(
                "/api/meetings",
                headers=headers,
                json={"title": "上传回显测试", "meeting_date": "2026-08-25"},
            )
            meeting_id = created.json()["id"]
            response = client.post(
                f"/api/meetings/{meeting_id}/sources",
                headers=headers,
                data={"text": "不应在响应中回显的识别稿正文"},
            )
            self.assertEqual(response.status_code, 201)
            payload = response.json()
            self.assertNotIn("text", payload)
            self.assertEqual(payload["text_chars"], 14)

    def test_duplicate_pasted_transcript_reuses_existing_source(self):
        headers = {"X-Meeting-Minutes-Action": "confirm"}
        with TestClient(app) as client:
            created = client.post(
                "/api/meetings", headers=headers,
                json={"title": "重复资料测试", "meeting_date": "2026-08-25"},
            )
            meeting_id = created.json()["id"]
            first = client.post(
                f"/api/meetings/{meeting_id}/sources", headers=headers,
                data={"text": "完全相同的识别稿", "source_name": "first.txt"},
            )
            second = client.post(
                f"/api/meetings/{meeting_id}/sources", headers=headers,
                data={"text": "完全相同的识别稿", "source_name": "second.txt"},
            )
            self.assertEqual(first.status_code, 201)
            self.assertEqual(second.status_code, 201)
            self.assertEqual(second.json()["id"], first.json()["id"])
            self.assertTrue(second.json()["duplicate_skipped"])
            sources = client.get(f"/api/meetings/{meeting_id}/sources").json()
            self.assertEqual(len(sources), 1)

    def test_date_is_required_before_organize_and_all_mutations_need_confirmation(self):
        headers = {"X-Meeting-Minutes-Action": "confirm"}
        with TestClient(app) as client:
            missing = client.post("/api/meetings", json={"title": "接口测试"})
            self.assertEqual(missing.status_code, 403)
            created = client.post(
                "/api/meetings", json={"title": "接口测试"}, headers=headers
            )
            self.assertEqual(created.status_code, 201)
            meeting_id = created.json()["id"]
            uploaded = client.post(
                f"/api/meetings/{meeting_id}/sources",
                data={"text": "说话人1：建议先开展测试。", "source_name": "识别稿.txt"},
                headers=headers,
            )
            self.assertEqual(uploaded.status_code, 201)
            organize = client.post(
                f"/api/meetings/{meeting_id}/organize", headers=headers
            )
            self.assertEqual(organize.status_code, 400)
            self.assertIn("会议日期", organize.json()["detail"])
            invalid = client.patch(
                f"/api/meetings/{meeting_id}",
                json={"title": "接口测试", "meeting_date": "2026/08/25"},
                headers=headers,
            )
            self.assertEqual(invalid.status_code, 400)
            updated = client.patch(
                f"/api/meetings/{meeting_id}",
                json={"title": "接口测试", "meeting_date": "2026-08-25"},
                headers=headers,
            )
            self.assertEqual(updated.status_code, 200)
            evil = client.post(
                "/api/meetings",
                json={"title": "跨站"},
                headers={**headers, "Origin": "http://attacker.example"},
            )
            self.assertEqual(evil.status_code, 403)
            cross_scheme = client.post(
                "/api/meetings",
                json={"title": "跨协议"},
                headers={**headers, "Origin": "https://testserver"},
            )
            self.assertEqual(cross_scheme.status_code, 403)


class StaticUiTests(unittest.TestCase):
    def test_complete_inline_script_has_valid_javascript_syntax(self):
        html = (Path(__file__).parents[1] / "app" / "static" / "index.html").read_text(
            encoding="utf-8"
        )
        match = re.search(r"<script>(.*)</script>", html, re.S)
        self.assertIsNotNone(match)
        node = shutil.which("node")
        if node is None:
            self.skipTest("node is unavailable for JavaScript syntax validation")
        checked = subprocess.run(
            [node, "--check"], input=match.group(1), text=True,
            capture_output=True, check=False, timeout=10,
        )
        self.assertEqual(checked.returncode, 0, checked.stderr)


def tearDownModule():
    connection = getattr(db._local, "conn", None)
    if connection is not None:
        connection.close()
        db._local.conn = None
    shutil.rmtree(_DATA_DIR, ignore_errors=True)


if __name__ == "__main__":
    unittest.main()
